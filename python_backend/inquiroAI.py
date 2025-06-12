"""
InquiroAI - Document Processing and Question Answering System
==========================================================

This module provides functionality for processing various types of documents (PDF, TXT, DOCX, images)
and answering questions about their content using local LLMs via Ollama.

Key Features:
- Multi-format document processing (PDF, TXT, DOCX, images)
- Text chunking and embedding generation
- Image description generation and embedding
- Question answering with context retrieval
- Document summarization
- Support for multiple LLM models

The system uses ChromaDB for vector storage and retrieval, and Langchain for document processing.

Sabareesh - 05/22/2025 - Version 1.0
"""

import os
import sys
import json
import uuid
import time
import logging
import tempfile
from typing import List, Dict, Optional, Any, Tuple, Union
import pyttsx3 # For Text-to-Speech
import whisper
import asyncio
from typing import Dict, Any, Optional

# Global variables
stt_model = None  # For Whisper model

def initialize_stt_model():
    """Initialize the Whisper speech-to-text model."""
    global stt_model
    if stt_model is None:
        try:
            logger.info("Loading Whisper STT model (this may take a moment on first run)...")
            stt_model = whisper.load_model("base")  # "base" is a good balance of accuracy and speed
            logger.info("Whisper STT model loaded successfully")
            return stt_model
        except Exception as e:
            logger.error(f"Error loading Whisper model: {e}", exc_info=True)
            return None
    return stt_model

from concurrent.futures import ThreadPoolExecutor

import chromadb
import fitz  # PyMuPDF
import ollama
from chromadb import Collection
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_chroma import Chroma
from langchain_community.vectorstores.utils import filter_complex_metadata
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
from langchain_core.vectorstores import VectorStoreRetriever

def get_user_config_dir() -> str:
    """Get the user's configuration directory for InquiroAI."""
    if sys.platform == 'win32':
        documents_dir = os.path.join(os.path.expanduser('~'), 'Documents')
    elif sys.platform == 'darwin':  # macOS
        documents_dir = os.path.join(os.path.expanduser('~'), 'Documents')
    else:  # Linux and other Unix-like
        documents_dir = os.path.join(os.path.expanduser('~'), 'Documents')
    
    config_dir = os.path.join(documents_dir, 'InquiroAI')
    os.makedirs(config_dir, exist_ok=True)
    return config_dir

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(threadName)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stderr)]
)
logger = logging.getLogger(__name__)

# --- Configuration ---
# Ensure paths are relative to this script's location or use absolute paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_DIR = get_user_config_dir()
MODELS_JSON_PATH = os.path.join(CONFIG_DIR, 'models.json')
CHROMA_PERSIST_DIR = os.path.join(CONFIG_DIR, "chroma_db_electron")
COLLECTION_NAME = "inquiroai_collection_v1"

# Processing Configuration
CHUNK_SIZE = 500  # Size of text chunks for splitting documents
CHUNK_OVERLAP = 100  # Overlap between chunks to maintain context
TEXT_RETRIEVER_K = 50  # Number of relevant chunks to retrieve for each query
IMAGE_RETRIEVER_K = 1  # Number of images to retrieve for each query

# Model Configuration - Initialize with None, will be set from models.json
TEXT_EMBEDDING_MODEL = None
IMAGE_EMBEDDING_MODEL = None
IMAGE_QUERY_MODEL = None
LLM_MODEL = None

# Enhanced error handling for models.json
try:
    if not os.path.exists(MODELS_JSON_PATH):
        default_models = {
            "model_categories": {
                "querying_models": ["llama2:13b"],
                "image_embedding_models": ["llava:13b"],
                "text_embedding_models": ["all-minilm:latest"]
            }
        }
        os.makedirs(os.path.dirname(MODELS_JSON_PATH), exist_ok=True)
        with open(MODELS_JSON_PATH, 'w') as f:
            json.dump(default_models, f, indent=4)
        logger.info(f"Created default models.json at {MODELS_JSON_PATH}")
        models_data = default_models
    else:
        with open(MODELS_JSON_PATH, 'r') as f:
            models_data = json.load(f)
            
    categories = models_data.get('model_categories', {})

    # Validate model categories
    if not all(isinstance(categories.get(cat, []), list) for cat in ['querying_models', 'image_embedding_models', 'text_embedding_models']):
        raise ValueError("Invalid models.json structure: Expected lists for all model categories.")

    # Set default models from the first available model in each category
    TEXT_EMBEDDING_MODEL = categories.get('text_embedding_models', [])[0] if categories.get('text_embedding_models') else "all-minilm:latest"
    IMAGE_EMBEDDING_MODEL = categories.get('image_embedding_models', [])[0] if categories.get('image_embedding_models') else "llava:13b"
    IMAGE_QUERY_MODEL = categories.get('querying_models', [])[0] if categories.get('querying_models') else "llama2:13b"
    LLM_MODEL = IMAGE_QUERY_MODEL  # Use the same model for text queries
    
except Exception as e:
    logger.error(f"Error loading models.json: {e}")
    # Set fallback defaults
    TEXT_EMBEDDING_MODEL = "all-minilm:latest"
    IMAGE_EMBEDDING_MODEL = "llava:13b"
    IMAGE_QUERY_MODEL = "llama2:13b"
    LLM_MODEL = IMAGE_QUERY_MODEL

# --- Global State ---
current_embedding_model: str = TEXT_EMBEDDING_MODEL
current_image_embedding_model: str = IMAGE_EMBEDDING_MODEL
current_querying_model: str = LLM_MODEL
stt_model = None # For Whisper STT model

# Async libraries
import asyncio
from functools import partial

# Dictionary to store state for each processed file
processed_files: Dict[str, Dict[str, Any]] = {}  # file_path -> state dict
# State will include:
# - processing_mode: str
# - retriever: VectorStoreRetriever
# - file_name: str
# Ensure the Chroma directory exists
os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)

# --- Helper Functions ---
def is_image_file(file_path: Optional[str]) -> bool:
    """
    Check if a file is an image based on its extension.
    
    Args:
        file_path: Path to the file to check
        
    Returns:
        bool: True if the file has an image extension, False otherwise
    """
    if not file_path:
        return False
    return os.path.splitext(file_path)[1].lower() in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]

def _get_chroma_collection(persist_dir: str = CHROMA_PERSIST_DIR, 
                         collection_name: str = COLLECTION_NAME) -> Collection:
    """
    Get or create a ChromaDB collection for storing embeddings.
    
    Args:
        persist_dir: Directory where ChromaDB data is stored
        collection_name: Name of the collection to get or create
        
    Returns:
        Collection: ChromaDB collection object
    """
    client = chromadb.PersistentClient(path=persist_dir)
    collection = client.get_or_create_collection(name=collection_name)
    return collection

def _generate_uuid(prefix: str = "") -> str:
    """
    Generate a unique identifier with an optional prefix.
    
    Args:
        prefix: Optional string prefix for the UUID
        
    Returns:
        str: Generated UUID with prefix if provided
    """
    return f"{prefix}_{uuid.uuid4()}"

def get_available_models() -> List[str]:
    """
    Get list of available models from models.json file.
    
    Returns:
        List[str]: List of available model names parsed from models.json categories
        
    Note:
        Models are now organized by categories in models.json under model_categories
    """
    try:
        if not os.path.exists(MODELS_JSON_PATH):
            logger.warning(f"models.json not found at {MODELS_JSON_PATH}, using default model list")
            default_models = ["llava-llama3:8b-v1.1-q4_0"]
            send_response({
                "type": "models_list",
                "models": default_models,
                "message": "Using default model list (models.json not found in user config directory)"
            })
            return default_models

        with open(MODELS_JSON_PATH, 'r') as f:
            models_data = json.load(f)
            
        # Extract all models from the model_categories
        available_models = []
        categories = models_data.get('model_categories', {})
        
        # Collect models from each category
        for category_name, category_models in categories.items():
            if isinstance(category_models, list):
                available_models.extend(category_models)
            
        if not available_models:
            logger.warning("No models found in models.json categories, using default model list")
            default_models = ["llava-llama3:8b-v1.1-q4_0"]
            send_response({
                "type": "models_list",
                "models": default_models,
                "message": "Using default model list (no models in models.json)"
            })
            return default_models

        # Remove duplicates while preserving order
        available_models = list(dict.fromkeys(available_models))
        
        logger.info(f"Loaded {len(available_models)} models from {len(categories)} categories in models.json")
        send_response({
            "type": "models_list",
            "models": available_models,
            "message": "Using models from models.json"
        })
        return available_models

    except Exception as e:
        logger.error(f"Error loading models from models.json: {e}", exc_info=True)
        default_models = ["llava-llama3:8b-v1.1-q4_0"]
        send_response({
            "type": "models_list",
            "models": default_models,
            "message": f"Error: {str(e)}. Using default model."
        })
        return default_models

def set_llm_model(model_name: str) -> str:
    """Set the current LLM model (simplified version)."""
    global current_llm_model
    try:
        available_models = get_models_by_category("querying_models")
        
        if model_name in available_models:
            current_llm_model = model_name
            logger.info(f"Set LLM model to: {model_name}")
            return f"Successfully set model to {model_name}"
        else:
            logger.warning(f"Requested model {model_name} not in available models. Using {current_llm_model}")
            return f"Model {model_name} not available. Using {current_llm_model}"
    except Exception as e:
        logger.error(f"Error setting LLM model: {e}", exc_info=True)
        return f"Error setting model: {str(e)}. Using {current_llm_model}"

def set_embedding_model(model_name: str) -> str:
    """
    Set the current embedding model for document processing.
    
    Args:
        model_name: Name of the model to use for embeddings
        
    Returns:
        str: Success/error message indicating the result
        
    Note:
        The model must be available in the models.json list
    """
    global current_embedding_model
    try:
        available_models = get_models_by_category("text_embedding_models")
        
        if model_name in available_models:
            current_embedding_model = model_name
            logger.info(f"Set embedding model to: {model_name}")
            return f"Successfully set embedding model to {model_name}"
        else:
            logger.warning(f"Requested model {model_name} not in available models. Using {current_embedding_model}")
            return f"Model {model_name} not available. Using {current_embedding_model}"
    except Exception as e:
        logger.error(f"Error setting embedding model: {e}", exc_info=True)
        return f"Error setting model: {str(e)}. Using {current_embedding_model}"

def set_querying_model(model_name: str) -> Dict[str, Any]:
    """
    Set the current querying model for answering questions.
    
    Args:
        model_name: Name of the model to use for querying
        
    Returns:
        Dict[str, Any]: Response dictionary with type and message
        
    Note:
        The model must be available in the models.json list
    """
    global current_querying_model
    try:
        available_models = get_models_by_category("querying_models")
        
        if model_name in available_models:
            current_querying_model = model_name
            logger.info(f"Set querying model to: {model_name}")
            return {
                "type": "model_set",
                "message": f"Successfully set querying model to {model_name}"
            }
        else:
            msg = f"Model {model_name} not available. Using {current_querying_model}"
            logger.warning(msg)
            return {
                "type": "error",
                "message": msg
            }
    except Exception as e:
        msg = f"Error setting model: {str(e)}. Using {current_querying_model}"
        logger.error(msg, exc_info=True)
        return {
            "type": "error",
            "message": msg
        }

def set_image_embedding_model(model_name: str) -> Dict[str, Any]:
    """
    Set the current image embedding model for processing images.
    
    Args:
        model_name: Name of the model to use for image embeddings
        
    Returns:
        Dict[str, Any]: Response dictionary with type and message
        
    Note:
        The model must be available in the models.json list
    """
    global current_image_embedding_model
    try:
        available_models = get_models_by_category("image_embedding_models")
        
        if model_name in available_models:
            current_image_embedding_model = model_name
            logger.info(f"Set image embedding model to: {model_name}")
            return {
                "type": "model_set",
                "message": f"Successfully set image embedding model to {model_name}"
            }
        else:
            msg = f"Model {model_name} not available. Using {current_image_embedding_model}"
            logger.warning(msg)
            return {
                "type": "error",
                "message": msg
            }
    except Exception as e:
        msg = f"Error setting model: {str(e)}. Using {current_image_embedding_model}"
        logger.error(msg, exc_info=True)
        return {
            "type": "error",
            "message": msg
        }

def get_models_by_category(category: str) -> List[str]:
    """
    Get list of models for a specific category from models.json.
    
    Args:
        category: Category name to get models for ('querying_models', 'image_embedding_models', 'text_embedding_models')
        
    Returns:
        List[str]: List of model names for the specified category
    """
    try:
        if not os.path.exists(MODELS_JSON_PATH):
            logger.warning(f"models.json not found at {MODELS_JSON_PATH}, using default models")
            if category == "querying_models":
                return ["llava-llama3:8b-v1.1-q4_0"]
            elif category == "image_embedding_models":
                return ["llava:13b"]
            else:  # text_embedding_models
                return ["all-minilm:latest"]

        # Read and parse models.json from user's config directory
        with open(MODELS_JSON_PATH, 'r') as f:
            models_data = json.load(f)
            
        categories = models_data.get('model_categories', {})
        models = categories.get(category, [])
        
        if not models:
            logger.warning(f"No models found for category {category}, using default")
            if category == "querying_models":
                return ["llava-llama3:8b-v1.1-q4_0"]
            elif category == "image_embedding_models":
                return ["llava:13b"]
            else:  # text_embedding_models
                return ["all-minilm:latest"]
                
        return models

    except Exception as e:
        logger.error(f"Error loading models for category {category}: {e}", exc_info=True)
        # Return appropriate default model based on category
        if category == "querying_models":
            return ["llava-llama3:8b-v1.1-q4_0"]
        elif category == "image_embedding_models":
            return ["llava:13b"]
        else:  # text_embedding_models
            return ["all-minilm:latest"]
# --- Core RAG Functions (mostly unchanged, minor logging adjustments) ---
# load_and_split_text_document, generate_text_embeddings_parallel,
# get_image_description, query_llava_with_image, query_text_llm,
# summarize_transcript
# (Keep these functions as they are in the original script, ensuring they use the logger)
# --- Placeholder for brevity - Assume these functions exist here ---
# --- (Copy the implementations from the original script) ---
# Example:
def load_and_split_text_document(file_path: str, chunk_size: int, chunk_overlap: int) -> Optional[List[Document]]:
    """Loads a TEXT-BASED document (PDF, TXT, DOCX), splits it, and cleans metadata."""
    logger.info(f"Loading text document: {file_path}")
    file_extension = os.path.splitext(file_path)[1].lower()
    documents: Optional[List[Document]] = None

    try:
        if file_extension == ".pdf":
            loader = PyMuPDFLoader(file_path)
            documents = loader.load()
        elif file_extension == ".txt":
            loader = TextLoader(file_path, encoding='utf-8', autodetect_encoding=True)
            documents = loader.load()
        elif file_extension == ".docx":
            try:
                from docx import Document as DocxDocument # Local import
                doc = DocxDocument(file_path)
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
                if text_content.strip():
                    documents = [Document(page_content=text_content, metadata={"source": os.path.basename(file_path)})]
                else:
                    documents = []
            except ImportError:
                logger.error("python-docx not installed. Run 'pip install python-docx'. Cannot process DOCX.")
                return None
            except Exception as e:
                logger.error(f"Error processing DOCX file {file_path}: {e}", exc_info=True)
                return None
        else:
            logger.error(f"Unsupported text document type: {file_extension}")
            return None

        if documents is None:
             logger.error(f"Loader failed unexpectedly for {file_path}")
             return None
        if not documents:
            logger.warning(f"No content loaded from {file_path}. File might be empty or unreadable.")
            return []

        logger.info(f"Loaded {len(documents)} initial document sections.")

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = text_splitter.split_documents(documents)
        cleaned_chunks = filter_complex_metadata(chunks)

        if not cleaned_chunks:
            logger.warning(f"Splitting resulted in zero valid chunks for {file_path}.")
            return []
        logger.info(f"Split into {len(cleaned_chunks)} chunks.")
        return cleaned_chunks

    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        return None
    except Exception as e:
        logger.error(f"Error loading/splitting document {file_path}: {e}", exc_info=True)
        return None

def generate_text_embeddings_parallel(chunks: List[Document], embed_func: OllamaEmbeddings) -> List[Tuple[int, List[float]]]:
    """Generates embeddings for text document chunks in parallel. Returns empty list on failure."""
    if not chunks or not embed_func:
        logger.warning("No text chunks or embedding function provided for embedding.")
        return []

    logger.info(f"Starting text embedding for {len(chunks)} chunks using '{embed_func.model}'...")
    results_with_indices: List[Tuple[int, Optional[List[float]]]] = []

    def _embed_chunk_with_index(index_chunk_tuple: Tuple[int, Document]) -> Tuple[int, Optional[List[float]]]:
        index, chunk = index_chunk_tuple
        try:
            content = getattr(chunk, 'page_content', '')
            if not content or content.isspace():
                logger.warning(f"Skipping embedding for chunk {index} due to empty content.")
                return index, None
            embedding = embed_func.embed_query(content)
            return index, embedding
        except Exception as e:
            logger.error(f"Error embedding chunk {index} with model '{embed_func.model}': {type(e).__name__}", exc_info=False)
            return index, None

    try:
        with ThreadPoolExecutor(thread_name_prefix="EmbedWorker") as executor:
            results_with_indices = list(executor.map(_embed_chunk_with_index, enumerate(chunks)))

        successful_embeddings = [(idx, emb) for idx, emb in results_with_indices if emb is not None]
        failed_count = len(chunks) - len(successful_embeddings)

        if failed_count > 0:
            logger.warning(f"{failed_count} text chunks failed during embedding.")
        if not successful_embeddings and chunks:
             logger.error(f"Text embedding failed for all {len(chunks)} input chunks.")
             return []

        logger.info(f"Successfully embedded {len(successful_embeddings)} text chunks.")
        successful_embeddings.sort(key=lambda x: x[0])
        return successful_embeddings
    except Exception as e:
        logger.error(f"Error during parallel text embedding execution: {e}", exc_info=True)
        return []

def get_image_description(image_path: str, model: str = IMAGE_QUERY_MODEL) -> Optional[str]:
    """Uses LLaVA via ollama.chat to generate a description of an image."""
    if not os.path.exists(image_path):
        logger.error(f"Image file not found for description generation: {image_path}")
        return None
    try:
        logger.info(f"Generating description for image: {os.path.basename(image_path)} using {model}")
        prompt = "Describe this image in detail, focusing on objects, scene, text, and overall context."
        response = ollama.chat(
            model=model,
            messages=[{'role': 'user', 'content': prompt, 'images': [image_path]}]
        )
        description = response.get('message', {}).get('content')
        if description:
            logger.info(f"Generated description for {os.path.basename(image_path)} (length: {len(description)})")
            return description.strip()
        else:
            logger.error(f"Failed to get description from LLaVA for {os.path.basename(image_path)}. Response: {response}")
            return None
    except ollama.ResponseError as ore:
        logger.error(f"Ollama API Error generating description for '{os.path.basename(image_path)}' with '{model}': {ore.status_code} - {ore.error}", exc_info=False)
        return None
    except Exception as e:
        logger.error(f"Error generating image description for {os.path.basename(image_path)}: {e}", exc_info=True)
        return None

def query_llava_with_image(question: str, image_path: str, llm_model: str = None) -> str:
    """Queries the specified Multimodal model (LLaVA) with a question and an image."""
    # Use IMAGE_QUERY_MODEL if no model provided. Don't use current_llm_model for image queries as it may not be multimodal.
    model_to_use = llm_model or IMAGE_QUERY_MODEL
    logger.info(f"Querying Multimodal model '{model_to_use}' with image '{os.path.basename(image_path)}' and question: '{question[:50]}...'")

    if not os.path.exists(image_path):
        logger.error(f"Image file not found at path for querying: {image_path}")
        return "Error: The image file specified could not be found for answering the question."

    try:
        start_time = time.time()
        response = ollama.chat(
            model=model_to_use,
            messages=[{'role': 'user', 'content': question, 'images': [image_path]}]
        )
        message_content = response.get('message', {}).get('content')
        if message_content:
            end_time = time.time()
            logger.info(f"Multimodal model '{model_to_use}' answered in {end_time - start_time:.2f} seconds.")
            return message_content.strip()
        else:
            logger.error(f"Multimodal model response missing content. Full response: {response}")
            return "Error: Received incomplete response from the multimodal model."
    except ollama.ResponseError as ore:
        logger.error(f"Ollama API Error (Multimodal Query): {ore.status_code} - {ore.error}", exc_info=False)
        return f"Error: Multimodal model API error ({ore.status_code}). Check Ollama server/model '{model_to_use}'. Details: {ore.error}"
    except Exception as e:
        logger.error(f"Error querying multimodal model '{model_to_use}': {e}", exc_info=True)
        return f"Error querying the multimodal model ({type(e).__name__})."

def query_text_llm(question: str, context: str, llm_model: str = None) -> str:
    """Formats prompt with TEXT context and queries the specified Ollama LLM."""
    # Use current_llm_model if no specific model is provided
    model_to_use = llm_model or current_llm_model
    logger.info(f"Querying Text LLM '{model_to_use}' for question: '{question[:50]}...'")

    if not context or context.startswith("Error:") or context == "No relevant context found.":
        context_issue = context if context else "Context is missing"
        logger.warning(f"Context indicates an issue or is missing: '{context_issue}'. Cannot answer based on document.")
        return f"I couldn't find relevant information in the document to answer your question. ({context_issue})"

    template = """You are an AI assistant. Answer the following question based *only* on the provided context.
If the context does not contain the answer, state "The provided context does not contain the answer to this question."
Do not use any external knowledge. Keep the answer concise.

Context:
---
{context}
---

Question: {question}

Answer:"""
    formatted_prompt = template.format(context=context, question=question)

    try:
        start_time = time.time()
        response = ollama.chat(model=model_to_use, messages=[{'role': 'user', 'content': formatted_prompt}])
        message_content = response.get('message', {}).get('content')
        if message_content:
            end_time = time.time()
            logger.info(f"LLM '{model_to_use}' answered in {end_time - start_time:.2f} seconds.")
            return message_content.strip()
        else:
            logger.error(f"Text LLM response missing content. Full response: {response}")            
            return "Error: Received incomplete response from the language model."
    except ollama.ResponseError as ore:
        logger.error(f"Ollama API Error (Text LLM): {ore.status_code} - {ore.error}", exc_info=False)
        return f"Error: Text LLM API error ({ore.status_code}). Check Ollama server/model '{model_to_use}'. Details: {ore.error}"
    except Exception as e:
        logger.error(f"Error querying Text LLM model '{model_to_use}': {e}", exc_info=True)
        return f"Error querying the text language model ({type(e).__name__})."

def summarize_transcript(transcript: str, max_chunk_chars: int = 2000, summary_model: str = LLM_MODEL) -> str:
    """Summarizes text using the LLM, handling long text by chunking."""
    if not transcript or transcript.isspace():
        return "No text provided to summarize."

    chunks = [transcript[i:i + max_chunk_chars] for i in range(0, len(transcript), max_chunk_chars)]
    if not chunks:
        logger.error("Failed to create chunks from non-empty transcript.")
        return "Error: Could not process the text into chunks for summarization."

    logger.info(f"Split transcript into {len(chunks)} chunks for summarization.")

    chunk_summaries = []
    for i, chunk in enumerate(chunks):
        prompt = f"Summarize the following text into concise key points (this is chunk {i+1} of {len(chunks)}):\n\nText Chunk:\n{chunk}\n\nSummary:"
        logger.info(f"Requesting summary for chunk {i+1}/{len(chunks)} using model {summary_model}")
        try:
            response = ollama.chat(model=summary_model, messages=[{'role': 'user', 'content': prompt}])
            content = response.get('message', {}).get('content')
            if content:
                 chunk_summaries.append(content.strip())
            else:
                 logger.warning(f"Received empty summary response for chunk {i+1}.")
                 chunk_summaries.append(f"[Summary for chunk {i+1} failed]")
        except Exception as e:
            logger.error(f"Error summarizing chunk {i+1} with model {summary_model}: {e}", exc_info=True)
            chunk_summaries.append(f"[Error summarizing chunk {i+1}]")

    combined_summary = "\n\n".join(chunk_summaries)

    if len(chunks) > 1 and len(combined_summary) > max_chunk_chars * 1.2:
        logger.info("Generating final meta-summary.")
        final_prompt = f"Synthesize the following chunk summaries into a single, coherent summary:\n\n{combined_summary}\n\nFinal Summary:"
        try:
             response = ollama.chat(model=summary_model, messages=[{'role': 'user', 'content': final_prompt}])
             content = response.get('message', {}).get('content')
             if content:
                  return content.strip()
             else:
                  logger.warning("Failed to generate final meta-summary. Returning combined chunk summaries.")
                  return combined_summary
        except Exception as e:
            logger.error(f"Error generating final meta-summary: {e}", exc_info=True)
            return combined_summary
    else:
        return combined_summary
# --- End Core Functions ---


# --- ChromaDB Management ---
_chroma_client = None

def initialize_chroma_client():
    """Initialize or reinitialize the ChromaDB client with proper settings."""
    global _chroma_client
    try:
        _chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
        # Try to get or create collection
        try:
            collection = _chroma_client.get_collection(COLLECTION_NAME)
            logger.info(f"Retrieved existing collection: {COLLECTION_NAME}")
        except Exception:
            collection = _chroma_client.create_collection(COLLECTION_NAME)
            logger.info(f"Created new collection: {COLLECTION_NAME}")
        return collection
    except Exception as e:
        logger.error(f"Error initializing ChromaDB client: {e}")
        raise

def reset_collection():
    """Reset the ChromaDB collection when changing embedding models."""
    global _chroma_client
    try:
        if _chroma_client is not None:
            try:
                _chroma_client.delete_collection(COLLECTION_NAME)
                logger.info(f"Deleted collection: {COLLECTION_NAME}")
            except Exception as e:
                logger.warning(f"Error deleting collection: {e}")
        
        # Reinitialize with new collection
        collection = initialize_chroma_client()
        logger.info("ChromaDB collection reset successfully")
        return collection
    except Exception as e:
        logger.error(f"Error resetting ChromaDB collection: {e}")
        raise

# --- Document Processing Helpers ---
def load_and_split_text_document(file_path: str) -> Optional[List[Document]]:
    """Loads a TEXT-BASED document (PDF, TXT, DOCX), splits it, and cleans metadata."""
    try:
        if file_path.lower().endswith('.pdf'):
            loader = PyMuPDFLoader(file_path)
        elif file_path.lower().endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8', autodetect_encoding=True)
        else:
            logger.error(f"Unsupported document type: {file_path}")
            return None

        documents = loader.load()
        if not documents:
            logger.warning(f"No content loaded from {file_path}")
            return []

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP
        )
        chunks = text_splitter.split_documents(documents)
        cleaned_chunks = filter_complex_metadata(chunks)

        if not cleaned_chunks:
            logger.warning(f"No valid chunks after processing {file_path}")
            return []

        logger.info(f"Split {file_path} into {len(cleaned_chunks)} chunks")
        return cleaned_chunks

    except Exception as e:
        logger.error(f"Error processing document {file_path}: {e}", exc_info=True)
        return None

async def load_multiple_documents_from_files(file_paths: List[str]) -> List[Document]:
    """
    Load and split multiple documents in parallel.
    
    Args:
        file_paths: List of file paths to process
    
    Returns:
        List of Document chunks from all files
    """
    all_chunks = []
    
    # Process each file
    for file_path in file_paths:
        try:
            chunks = load_and_split_text_document(file_path)
            if chunks:
                all_chunks.extend(chunks)
                logger.info(f"Successfully processed {file_path} into {len(chunks)} chunks")
            else:
                logger.warning(f"No chunks extracted from {file_path}")
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            continue
            
    return all_chunks

async def process_text_document(file_path: str) -> List[Document]:
    """Process a text document and return chunks."""
    try:
        if file_path.lower().endswith('.pdf'):
            loader = PyMuPDFLoader(file_path)
        else:  # Default to text loader for .txt and other text files
            loader = TextLoader(file_path, encoding='utf-8')
            
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP
        )
        chunks = text_splitter.split_documents(documents)
        return chunks
    except Exception as e:
        logger.error(f"Error processing text document: {e}")
        raise

async def generate_embedding(text: str, embedding_model) -> List[float]:
    """Generate embedding for a piece of text."""
    try:
        embedding = await embedding_model.aembed_query(text)
        return embedding
    except Exception as e:
        logger.error(f"Error generating embedding: {e}")
        raise

# --- Model Management ---
async def set_embedding_model(model_name: str) -> Dict[str, Any]:
    """Set the text embedding model and reset the collection to handle dimension changes."""
    global current_embedding_model, _chroma_client
    try:
        # Reset collection when changing embedding models
        collection = reset_collection()
        
        current_embedding_model = model_name
        logger.info(f"Set text embedding model to: {model_name}")
        return {"type": "model_set", "message": f"Text embedding model set to {model_name}"}
    except Exception as e:
        logger.error(f"Error setting embedding model: {e}")
        return {"type": "error", "message": f"Failed to set embedding model: {str(e)}"}

# --- Document Processing ---
async def process_document_internal(file_path: str, mode: str = 'text') -> Union[Dict[str, Any], None]:
    """Internal document processing function with improved error handling."""
    try:
        # Initialize embeddings with current model
        embeddings = OllamaEmbeddings(
            model=current_embedding_model if mode == 'text' else current_image_embedding_model
        )
        
        # Get or create ChromaDB collection
        collection = initialize_chroma_client()
        
        # Process the document based on mode
        if mode == 'text':
            # Process text document
            chunks = await process_text_document(file_path)
            if not chunks:
                raise ValueError("No text chunks extracted from document")
                
            # Create embeddings in batches
            ids_to_add = []
            embeddings_to_add = []
            metadatas_to_add = []
            documents_to_add = []
            
            for i, chunk in enumerate(chunks):
                try:
                    # Generate embedding
                    embedding = await generate_embedding(chunk.page_content, embeddings)
                    
                    ids_to_add.append(f"{os.path.basename(file_path)}_{i}")
                    embeddings_to_add.append(embedding)
                    metadatas_to_add.append({"source": file_path, "chunk": i})
                    documents_to_add.append(chunk.page_content)
                    
                except Exception as e:
                    logger.error(f"Error processing chunk {i}: {e}")
                    continue
            
            if not embeddings_to_add:
                raise ValueError("No valid embeddings generated")
                
            # Add to collection
            try:
                collection.add(
                    ids=ids_to_add,
                    embeddings=embeddings_to_add,
                    metadatas=metadatas_to_add,
                    documents=documents_to_add
                )
            except Exception as e:
                if "dimension" in str(e).lower():
                    # Handle dimension mismatch
                    logger.warning("Embedding dimension mismatch detected. Resetting collection...")
                    collection = reset_collection()
                    # Try adding again with fresh collection
                    collection.add(
                        ids=ids_to_add,
                        embeddings=embeddings_to_add,
                        metadatas=metadatas_to_add,
                        documents=documents_to_add
                    )
                else:
                    raise
                    
            return {"status": "success", "chunks_processed": len(chunks)}
            
    except Exception as e:
        logger.critical(f"Unhandled error during document processing for {os.path.basename(file_path)}: {str(e)}", exc_info=True)
        return {"status": "error", "message": str(e)}

async def handle_process_document(file_path: str) -> Dict[str, Any]:
    """Process a document and store its embeddings."""
    try:
        if not os.path.exists(file_path):
            return {"type": "error", "message": f"File not found: {file_path}"}

        # Create a temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                result = await process_document_internal(file_path)
                if result["status"] == "success":
                    processing_mode = "text"  # We'll enhance this later for different modes
                    status = f"Processed {os.path.basename(file_path)} - {result['chunks_processed']} chunks"
                    return {
                        "type": "process_complete",
                        "status": status,
                        "processing_mode": processing_mode
                    }
                else:
                    raise Exception(result["message"])
                    
            except Exception as e:
                logger.error(f"Error processing document: {e}")
                return {
                    "type": "error",
                    "message": f"Processing failed: {str(e)}"
                }
            finally:
                logger.info(f"Cleaned up temp directory: {temp_dir}")

    except Exception as e:
        logger.error(f"Unhandled error in process_document: {e}")
        return {
            "type": "error",
            "message": f"Unhandled error: {str(e)}"
        }

# --- Async Document Processing ---
async def process_document(request_data: Dict[str, Any]) -> Dict[str, Any]:
    """Main document processing handler."""
    file_path = request_data.get('file_path')
    request_id = request_data.get('request_id')
    
    if not file_path or not os.path.exists(file_path):
        return {
            "type": "error",
            "message": f"Invalid or missing file: {file_path}",
            "request_id": request_id
        }

    try:
        # Process document and get result
        result = await handle_document_processing(file_path)
        result["request_id"] = request_id
        return result
    except Exception as e:
        logger.error(f"Error in process_document: {e}", exc_info=True)
        return {
            "type": "error",
            "message": f"Processing failed: {str(e)}",
            "request_id": request_id
        }

async def handle_document_processing(file_path: str) -> Dict[str, Any]:
    """Process a document with async support."""
    try:
        # Create temporary processing directory
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                logger.info(f"Processing '{os.path.basename(file_path)}'...")
                status_messages = []
                
                # For text documents (txt, pdf, docx)
                if file_path.lower().endswith(('.txt', '.pdf', '.docx')):
                    status_messages.append('Loading and splitting document...')
                    
                    # Load and process document
                    chunks = await process_text_document(file_path)
                    if not chunks:
                        raise ValueError(f"No valid chunks extracted from {file_path}")
                    
                    status_messages.append(f"Successfully split into {len(chunks)} chunks")
                    
                    # Initialize embeddings
                    embeddings = OllamaEmbeddings(model=current_embedding_model)
                    
                    # Get collection
                    collection = initialize_chroma_client()
                    
                    # Generate embeddings in parallel
                    successful_embeddings = generate_text_embeddings_parallel(chunks, embeddings)
                    if not successful_embeddings:
                        raise ValueError("Failed to generate embeddings")
                    
                    # Prepare data for collection
                    ids = [f"{os.path.basename(file_path)}_{i}" for i, _ in successful_embeddings]
                    embeddings_list = [emb for _, emb in successful_embeddings]
                    metadatas = [{"source": file_path, "chunk": i} for i, _ in enumerate(successful_embeddings)]
                    documents = [chunks[i].page_content for i, _ in enumerate(successful_embeddings)]

                    # Add to collection
                    try:
                        collection.add(
                            ids=ids,
                            embeddings=embeddings_list,
                            metadatas=metadatas,
                            documents=documents
                        )
                    except Exception as e:
                        if "dimension" in str(e).lower():
                            collection = reset_collection()
                            collection.add(
                                ids=ids,
                                embeddings=embeddings_list,
                                metadatas=metadatas,
                                documents=documents
                            )
                        else:
                            raise

                    status = "\n".join(status_messages + [
                        f"Generated embeddings for {len(successful_embeddings)} chunks",
                        "Stored in vector database",
                        f"Processing complete for {os.path.basename(file_path)}"
                    ])
                    return {
                        "type": "process_complete",
                        "status": status,
                        "chunks_processed": len(successful_embeddings)
                    }
                else:
                    raise ValueError(f"Unsupported file type: {file_path}")
                    
            except Exception as e:
                status = "\n".join(status_messages + [f"Critical Error: {type(e).__name__}. Check logs."])
                logger.error(f"Processing failed: {e}", exc_info=True)
                return {
                    "type": "error",
                    "message": f"Processing failed: {status}",
                }
            finally:
                logger.info(f"Cleaned up temp directory: {temp_dir}")
    except Exception as e:
        logger.error(f"Unhandled error in handle_document_processing: {e}", exc_info=True)
        return {
            "type": "error",
            "message": f"Unhandled error: {str(e)}"
        }

# --- Utility Functions ---
def run_async(func):
    """Decorator to run async functions in the event loop."""
    def wrapper(*args, **kwargs):
        return asyncio.get_event_loop().run_until_complete(func(*args, **kwargs))
    return wrapper

# --- Response Processing ---
async def process_response_internal(query_text: str) -> Dict[str, Any]:
    """Process a query and generate a response."""
    try:
        result = await process_query(query_text)
        
        # Make sure we always return a query_result type
        if result.get('type') == 'query_response':
            result['type'] = 'query_result'
            
        return result
    except Exception as e:
        logger.error(f"Error processing response: {e}", exc_info=True)
        return {
            "type": "query_result",
            "success": False,
            "message": f"Error processing query: {str(e)}",
            "answer": None,
            "sources": [],
            "chunks_used": 0
        }
        
    except Exception as e:
        logger.error(f"Error processing response: {e}", exc_info=True)
        raise

async def process_query(query_text: str, collection_name: str = COLLECTION_NAME) -> Dict[str, Any]:
    """
    Process a query against the document collection and return the response.
    
    Args:
        query_text (str): The user's query
        collection_name (str): Name of the ChromaDB collection to query
        
    Returns:
        Dict[str, Any]: Response containing the answer and status
    """
    try:
        logger.info(f"Processing query: {query_text}")
        
        # Get collection and ensure it has documents
        collection = initialize_chroma_client()
        if collection.count() == 0:
            no_docs_message = "No documents have been processed yet. Please process documents first."
            audio_path = text_to_speech(no_docs_message)
            return {
                "type": "query_result",
                "success": False,
                "message": no_docs_message,
                "answer": None,
                "sources": [],
                "chunks_used": 0,
                "audio_response_path": audio_path
            }
        
        # Initialize query embeddings
        embeddings = OllamaEmbeddings(model=current_embedding_model)
        query_embedding = await embeddings.aembed_query(query_text)
        
        # Get relevant chunks
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=TEXT_RETRIEVER_K,
            include=['documents', 'metadatas']
        )
        
        if not results or not results['documents']:
            no_results_message = "No relevant context found in the processed documents."
            audio_path = text_to_speech(no_results_message)
            return {
                "type": "query_result",
                "success": False,
                "message": no_results_message,
                "answer": None,
                "sources": [],
                "chunks_used": 0,
                "audio_response_path": audio_path
            }
            
        # Build context from chunks with safer metadata handling
        contexts = []
        sources = set()
        for doc, meta in zip(results['documents'][0], results['metadatas'][0]):
            # Handle potential None metadata or missing source field
            source = "Unknown Source"
            if meta and isinstance(meta, dict) and 'source' in meta:
                source = os.path.basename(meta['source'])
            sources.add(source)
            contexts.append(f"From {source}:\n{doc}")
        
        context = "\n\n".join(contexts)
        
        # Format prompt with specific instruction to use only provided context
        prompt = f"""Using *only* the context provided below, answer the question.  
If the answer cannot be determined from the context, say so.

Context:  
{context}

Question: {query_text}

Answer:"""

        # Get answer from LLM using chat API with error handling
        try:
            if not current_querying_model:
                raise ValueError("No querying model selected")
                
            response = ollama.chat(
                model=current_querying_model,
                messages=[{"role": "user", "content": prompt}]
            )
            
            logger.debug(f"Raw LLM response: {response}")
            
            # Extract content from the response, handling different response formats
            answer = None
            if hasattr(response, 'message'):
                # Object-style response
                if hasattr(response.message, 'content'):
                    answer = response.message.content
                elif isinstance(response.message, dict) and 'content' in response.message:
                    answer = response.message['content']
            elif isinstance(response, dict):
                # Dictionary-style response
                message = response.get('message')
                if isinstance(message, dict):
                    answer = message.get('content')
                elif hasattr(message, 'content'):
                    answer = message.content
            
            if not answer:
                logger.error(f"Could not extract answer from response: {response}")
                raise ValueError("No valid answer content in LLM response")
                
        except Exception as llm_error:
            error_message = f"Error getting answer from LLM: {str(llm_error)}"
            logger.error(error_message, exc_info=True)
            audio_path = text_to_speech(error_message)
            return {
                "type": "query_result",
                "success": False,
                "message": error_message,
                "answer": None,
                "sources": list(sources),
                "chunks_used": len(contexts),
                "audio_response_path": audio_path
            }
        
        # Generate audio response with retry logic
        audio_path = None
        max_retries = 3
        retry_delay = 1  # seconds
        
        for attempt in range(max_retries):
            try:
                audio_path = text_to_speech(answer)
                if audio_path and os.path.exists(audio_path) and os.path.getsize(audio_path) > 0:
                    break
                if attempt < max_retries - 1:
                    await asyncio.sleep(retry_delay)
            except Exception as e:
                logger.warning(f"TTS attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    await asyncio.sleep(retry_delay)
                else:
                    logger.error("All TTS attempts failed")
        
        result = {
            "type": "query_result",
            "success": True,
            "message": "Query processed successfully",
            "answer": answer,
            "sources": list(sources),
            "chunks_used": len(contexts),
            "audio_response_path": audio_path
        }
        
        # Log if audio wasn't generated
        if not audio_path:
            logger.warning("Audio response not available for successful query")
            
        return result
        
    except Exception as e:
        error_message = f"Error processing query: {str(e)}"
        try:
            audio_path = text_to_speech(error_message)
        except Exception as tts_error:
            logger.error(f"Failed to create audio for error message: {tts_error}")
            audio_path = None
            
        logger.error(f"Error processing query: {e}", exc_info=True)
        return {
            "type": "query_result",
            "success": False,
            "message": error_message,
            "answer": None,
            "sources": [],
            "chunks_used": 0,
            "audio_response_path": audio_path
        }
# --- Main Request Handler ---
def handle_request(request_data: Dict[str, Any]) -> Dict[str, Any]:
    """Main request handler with async support."""
    request_type = request_data.get('type')
    
    try:
        if request_type == 'process_file':
            return run_async(process_document)(request_data)
        elif request_type == 'set_embedding_model':
            return run_async(set_embedding_model)(request_data.get('model_name'))
        elif request_type == 'submit_query':
            return run_async(lambda: process_response_internal(request_data.get('query')))()
        elif request_type == 'process_query':
            return run_async(lambda: process_query(request_data.get('query')))()
        else:
            return {
                "type": "error",
                "message": f"Unknown request type: {request_type}",
                "request_id": request_data.get('request_id')
            }
    except Exception as e:
        logger.error(f"Error in handle_request: {e}", exc_info=True)
        return {
            "type": "error",
            "message": str(e),
            "request_id": request_data.get('request_id')
        }
# --- Stdin/Stdout Communication ---

def send_response(data: Dict[str, Any]):
    """Send a JSON response to stdout for Electron to read."""
    try:
        # Ensure the data has all required fields
        if 'type' not in data:
            data['type'] = 'error'
            data['message'] = 'Response missing type field'
        
        response_json = json.dumps(data)
        print(response_json, flush=True)  # flush=True is critical for IPC
    except Exception as e:
        error_json = json.dumps({
            'type': 'error',
            'message': f'Error sending response: {str(e)}'
        })
        print(error_json, flush=True)

# --- Main Loop ---
async def process_files(file_paths: List[str]) -> Dict[str, Any]:
    """
    Process multiple files and store their embeddings in the vector store.
    
    Args:
        file_paths: List of file paths to process
        
    Returns:
        Dict with processing status and message
    """
    try:
        logger.info(f"Processing files: {file_paths}")
        
        # Validate file paths
        valid_paths = []
        for path in file_paths:
            if os.path.exists(path):
                valid_paths.append(path)
            else:
                logger.warning(f"File not found: {path}")
                
        if not valid_paths:
            return {
                "type": "processing_status",
                "success": False,
                "message": "No valid files found at the provided paths."
            }
        
        # Initialize ChromaDB client
        client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
        
        try:
            collection = client.get_collection(COLLECTION_NAME)
            logger.info("Got existing collection")
        except:
            collection = client.create_collection(COLLECTION_NAME)
            logger.info("Created new collection")
            
        # Initialize embeddings model
        try:
            embed_model = OllamaEmbeddings(
                model=current_embedding_model or TEXT_EMBEDDING_MODEL or "all-minilm"
            )
            logger.info(f"Initialized embedding model: {embed_model.model}")
        except Exception as e:
            logger.error(f"Failed to initialize embedding model: {e}")
            return {
                "type": "processing_status",
                "success": False,
                "message": f"Failed to initialize embedding model: {str(e)}"
            }
            
        # Load and split documents
        all_chunks = await load_multiple_documents_from_files(valid_paths)
        if not all_chunks:
            return {
                "type": "processing_status",
                "success": False,
                "message": "No valid content found in the provided files."
            }
            
        # Generate embeddings asynchronously
        embeddings = []
        ids = []
        texts = []
        
        # Process chunks in batches
        batch_size = 10
        total_chunks = len(all_chunks)
        processed = 0
        
        for i in range(0, total_chunks, batch_size):
            batch = all_chunks[i:i + batch_size]
            batch_embeddings = []
            
            try:
                # Generate embeddings for batch
                for chunk in batch:
                    try:
                        embedding = await generate_embedding(chunk.page_content, embed_model)
                        batch_embeddings.append((chunk, embedding))
                    except Exception as e:
                        logger.error(f"Failed to generate embedding for chunk: {e}")
                        continue
                
                # Add successful embeddings to lists
                for chunk, embedding in batch_embeddings:
                    embeddings.append(embedding)
                    ids.append(f"doc_{uuid.uuid4()}")
                    texts.append(chunk.page_content)
                
                processed += len(batch_embeddings)
                logger.info(f"Processed {processed}/{total_chunks} chunks")
                
            except Exception as batch_error:
                logger.error(f"Error processing batch: {batch_error}")
                continue
        
        if not embeddings:
            return {
                "type": "processing_status",
                "success": False,
                "message": "Failed to generate any valid embeddings."
            }
            
        try:
            # Add to collection
            collection.add(
                embeddings=embeddings,
                documents=texts,
                ids=ids
            )
            
            message = f"Successfully processed {len(valid_paths)} files.\n"
            message += f"Generated embeddings for {len(embeddings)} chunks."
            
            return {
                "type": "processing_status",
                "success": True,
                "message": message
            }
            
        except Exception as add_error:
            logger.error(f"Failed to add embeddings to collection: {add_error}")
            return {
                "type": "processing_status",
                "success": False,
                "message": f"Failed to store embeddings: {str(add_error)}"
            }
        
    except Exception as e:
        logger.error(f"Error processing files: {e}", exc_info=True)
        return {
            "type": "processing_status",
            "success": False,
            "message": f"Error processing files: {str(e)}"
        }

async def handle_command(command_data: Dict[str, Any]):
    """Handle individual commands asynchronously."""
    command = command_data.get('type')
    request_id = command_data.get('request_id', 'unknown')
    
    logger.info(f"Processing command: {command} (request_id: {request_id})")
    
    try:
        if command == 'get_models_by_category':
            category = command_data.get('category')
            if not category:
                return {
                    'type': 'error',
                    'message': 'No category provided',
                    'request_id': request_id
                }
            models = get_models_by_category(category)
            return {
                'type': 'models_list',
                'models': models,
                'request_id': request_id
            }
        
        elif command == 'process_files':
            file_paths = command_data.get('file_paths', [])
            if not file_paths:
                return {
                    'type': 'error',
                    'message': 'No file paths provided',
                    'request_id': request_id
                }

            result = await process_files(file_paths)
            result['request_id'] = request_id
            return result
                
        elif command == 'query':
            query_text = command_data.get('query_text') or command_data.get('query')  # Support both formats
            if not query_text:
                return {
                    'type': 'error',
                    'message': 'No query text provided',
                    'request_id': request_id
                }

            logger.info(f"Processing query: '{query_text}'")
            result = await process_query(query_text)
            result['request_id'] = request_id
            logger.info(f"Query result: {result}")
            return result
            
        elif command == 'get_models':
            models = get_available_models()
            return {
                'type': 'models_list',
                'models': models,
                'request_id': request_id,
                'message': 'Retrieved available models'
            }
            
        # Handle model setting commands
        elif command in ['set_embedding_model', 'set_image_embedding_model', 'set_querying_model']:
            model_name = command_data.get('model_name')
            if not model_name:
                return {
                    'type': 'error',
                    'message': 'No model name provided',
                    'request_id': request_id
                }
            
            logger.info(f"Setting {command.replace('set_', '')} to {model_name}")
            
            # Map command to corresponding function
            model_setters = {
                'set_embedding_model': set_embedding_model,
                'set_image_embedding_model': set_image_embedding_model,
                'set_querying_model': set_querying_model
            }
            
            try:
                setter_func = model_setters.get(command)
                if not setter_func:
                    raise ValueError(f"Invalid model setting command: {command}")
                
                # Call the setter function - note that set_embedding_model is async
                result = await setter_func(model_name) if command == 'set_embedding_model' else setter_func(model_name)
                
                # Convert string responses to dict format
                if isinstance(result, str):
                    if 'Successfully' in result:
                        result = {'type': 'model_set', 'message': result}
                    else:
                        result = {'type': 'error', 'message': result}
                        
                result['request_id'] = request_id
                logger.info(f"Model set result: {result}")
                return result
                
            except Exception as e:
                error_msg = f"Error setting {command.replace('set_', '')}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                return {
                    'type': 'error',
                    'message': error_msg,
                    'request_id': request_id
                }
        
        elif command == 'transcribe_audio':
            audio_path = command_data.get('audio_path')
            if not audio_path:
                return {'type': 'error', 'message': 'No audio path provided for transcription', 'request_id': request_id}
            
            result = await speech_to_text_async(audio_path)
            result['request_id'] = request_id
            return result
            
        else:
            logger.warning(f"Unknown command received: {command}")
            return {
                'type': 'error',
                'message': f'Unknown command: {command}',
                'request_id': request_id
            }
    except Exception as e:
        logger.error(f"Error processing command: {e}", exc_info=True)
        return {
            'type': 'error',
            'message': f'Command processing error: {str(e)}',
            'request_id': request_id
        }

def text_to_speech(text_to_speak: str) -> Optional[str]:
    """
    Converts text to speech using pyttsx3 and saves it as a temporary WAV file.
    Returns the path to the audio file, or None if failed.
    """
    logger.info(f"TTS: Attempting to speak: '{text_to_speak[:100]}...'")
    if not text_to_speak or not text_to_speak.strip():
        logger.warning("TTS: No text provided to speak.")
        return None

    output_filepath = None
    engine = None
    try:
        # Initialize the TTS engine
        engine = pyttsx3.init()
        
        # Optional: Configure voice properties for better quality
        engine.setProperty('rate', 150)  # Speed of speech
        engine.setProperty('volume', 0.9)  # Volume level
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio_file:
            output_filepath = temp_audio_file.name
        
        logger.info(f"TTS: Saving speech to temporary file: {output_filepath}")
        engine.save_to_file(text_to_speak, output_filepath)
        engine.runAndWait()

        # Verify the file was created successfully
        if os.path.exists(output_filepath) and os.path.getsize(output_filepath) > 0:
            logger.info(f"TTS: Speech successfully saved to {output_filepath}")
            return output_filepath
        else:
            logger.error(f"TTS: Failed to create speech file or file is empty: {output_filepath}")
            if os.path.exists(output_filepath):
                os.remove(output_filepath)
            return None

    except Exception as e:
        logger.error(f"TTS: Error during text-to-speech conversion: {e}", exc_info=True)
        # Clean up the temporary file if it exists
        if output_filepath and os.path.exists(output_filepath):
            try:
                os.remove(output_filepath)
            except Exception as cleanup_error:
                logger.warning(f"TTS: Failed to clean up temporary file: {cleanup_error}")
        return None

    finally:
        # Always try to properly stop the TTS engine
        if engine:
            try:
                engine.stop()
            except Exception as stop_error:
                logger.warning(f"TTS: Error stopping TTS engine: {stop_error}")

def initialize_stt_model(model_name="base"):
    """Initializes the Whisper STT model if not already loaded."""
    global stt_model
    if stt_model is None:
        logger.info(f"Loading Whisper STT model ({model_name}). This might take a moment...")
        try:
            stt_model = whisper.load_model(model_name)
            logger.info("Whisper STT model loaded successfully.")
        except Exception as e:
            logger.error(f"Error loading Whisper STT model: {e}", exc_info=True)
            stt_model = None # Ensure it's None if loading failed
    return stt_model

async def speech_to_text_async(audio_filepath: str) -> Dict[str, Any]:
    """
    Transcribes audio from a file path using Whisper.
    Runs the synchronous whisper.transcribe in a thread pool executor.
    Returns a dictionary with type and either text or error message.
    """
    model = initialize_stt_model()  # Ensure model is loaded
    if not model:
        return {"type": "error", "message": "Speech-to-text model is not available."}

    if not os.path.exists(audio_filepath):
        logger.error(f"Audio file not found for STT: {audio_filepath}")
        return {"type": "error", "message": f"Audio file not found at path: {audio_filepath}"}

    try:
        logger.info(f"Starting transcription for audio file: {audio_filepath}")
        loop = asyncio.get_event_loop()
        # Whisper's transcribe is CPU-bound, so run_in_executor is appropriate
        # Pass options as part of the second argument
        result = await loop.run_in_executor(None, lambda: model.transcribe(audio_filepath, fp16=False))
        transcribed_text = result.get("text", "")
        
        if not transcribed_text:
            return {"type": "error", "message": "Transcription produced no text"}
            
        logger.info(f"Transcription successful. Text (first 100 chars): '{transcribed_text[:100]}...'")
        return {"type": "transcription_result", "text": transcribed_text.strip()}
    except Exception as e:
        logger.error(f"Error during speech transcription for {audio_filepath}: {e}", exc_info=True)
        return {"type": "error", "message": f"Transcription failed: {str(e)}"}

def main_loop():
    """Main loop to handle requests from Electron."""
    logger.info("Python backend started")
    
    # Initialize asyncio event loop
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)

    try:
        # Initialize models and categories
        logger.info("Initializing model categories...")
        for category in ['querying_models', 'image_embedding_models', 'text_embedding_models']:
            models = get_models_by_category(category)
            logger.info(f"Loaded {len(models)} models for category {category}: {models}")

        for line in sys.stdin:
            try:
                command_data = json.loads(line)
                response = loop.run_until_complete(handle_command(command_data))
                send_response(response)
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse command JSON: {e}")
                send_response({
                    'type': 'error',
                    'message': f'Invalid JSON: {str(e)}',
                    'request_id': 'unknown'
                })
            except Exception as e:
                logger.error(f"Error processing command: {e}", exc_info=True)
                send_response({
                    'type': 'error',
                    'message': f'Command processing error: {str(e)}',
                    'request_id': 'unknown'
                })
    except Exception as e:
        logger.critical(f"Critical error in main loop: {e}", exc_info=True)
        sys.exit(1)
    finally:
        loop.close()

if __name__ == "__main__":
    main_loop()

def validate_models_json() -> None:
    """Validate models.json exists and has correct structure, create with defaults if missing."""
    default_models = {
        "model_categories": {
            "querying_models": [
                "llama2:13b",
                "llama3.1:8b",
                "deepseek-r1:7b"
            ],
            "image_embedding_models": [
                "llava:13b",
                "llava-llama3:8b-v1.1-q4_0"
            ],
            "text_embedding_models": [
                "all-minilm:latest",
                "nomic-embed-text:latest"
            ]
        }
    }

    try:
        if not os.path.exists(MODELS_JSON_PATH):
            logger.info(f"models.json not found at {MODELS_JSON_PATH}, creating with defaults...")
            os.makedirs(os.path.dirname(MODELS_JSON_PATH), exist_ok=True)
            with open(MODELS_JSON_PATH, 'w') as f:
                json.dump(default_models, f, indent=4)
            logger.info("Created default models.json")
            return

        # Validate existing file
        with open(MODELS_JSON_PATH, 'r') as f:
            models_data = json.load(f)
        
        categories = models_data.get('model_categories', {})
        required_categories = ['querying_models', 'image_embedding_models', 'text_embedding_models']
        
        # Check if all required categories exist and are lists
        missing_categories = []
        invalid_categories = []
        for category in required_categories:
            if category not in categories:
                missing_categories.append(category)
            elif not isinstance(categories[category], list):
                invalid_categories.append(category)
        
        if missing_categories or invalid_categories:
            logger.warning(f"Invalid models.json structure detected")
            if missing_categories:
                logger.warning(f"Missing categories: {missing_categories}")
            if invalid_categories:
                logger.warning(f"Invalid category types: {invalid_categories}")
            
            # Merge with defaults for missing/invalid categories
            for category in missing_categories + invalid_categories:
                categories[category] = default_models['model_categories'][category]
            
            # Save corrected file
            models_data['model_categories'] = categories
            with open(MODELS_JSON_PATH, 'w') as f:
                json.dump(models_data, f, indent=4)
            logger.info("Fixed and saved models.json with correct structure")
            
    except Exception as e:
        logger.error(f"Error handling models.json: {e}")
        logger.info("Creating new models.json with defaults")
        with open(MODELS_JSON_PATH, 'w') as f:
            json.dump(default_models, f, indent=4)

# Call validation on startup
validate_models_json()

async def load_multiple_documents_from_files(file_paths: List[str], chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> List[Document]:
    """
    Load and split multiple documents in parallel.
    
    Args:
        file_paths: List of file paths to process
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of Document chunks from all files
    """
    all_chunks = []
    
    # Process each file
    for file_path in file_paths:
        try:
            chunks = await process_text_document(file_path)
            if chunks:
                all_chunks.extend(chunks)
                logger.info(f"Successfully processed {file_path} into {len(chunks)} chunks")
            else:
                logger.warning(f"No chunks extracted from {file_path}")
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            continue
            
    return all_chunks

async def process_query(query_text: str) -> Dict[str, Any]:
    """
    Process a query using RAG approach with context from all available document chunks.
    Generates both text and audio responses.
    
    Args:
        query_text: The user's question or 'summarize' request
        
    Returns:
        Dict containing answer, audio response path, and metadata about sources used
    """
    try:
        # Get collection and ensure it has documents
        collection = initialize_chroma_client()
        if collection.count() == 0:
            no_docs_message = "No documents have been processed yet. Please process documents first."
            audio_path = text_to_speech(no_docs_message)
            return {
                "type": "query_result",
                "success": False,
                "message": no_docs_message,
                "answer": None,
                "sources": [],
                "chunks_used": 0,
                "audio_response_path": audio_path
            }
        
        # Initialize query embeddings
        embeddings = OllamaEmbeddings(model=current_embedding_model)
        query_embedding = await embeddings.aembed_query(query_text)
        
        # Get relevant chunks
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=TEXT_RETRIEVER_K,
            include=['documents', 'metadatas']
        )
        
        if not results or not results['documents']:
            no_results_message = "No relevant context found in the processed documents."
            audio_path = text_to_speech(no_results_message)
            return {
                "type": "query_result",
                "success": False,
                "message": no_results_message,
                "answer": None,
                "sources": [],
                "chunks_used": 0,
                "audio_response_path": audio_path
            }
            
        # Build context from chunks with safer metadata handling
        contexts = []
        sources = set()
        for doc, meta in zip(results['documents'][0], results['metadatas'][0]):
            # Handle potential None metadata or missing source field
            source = "Unknown Source"
            if meta and isinstance(meta, dict) and 'source' in meta:
                source = os.path.basename(meta['source'])
            sources.add(source)
            contexts.append(f"From {source}:\n{doc}")
        
        context = "\n\n".join(contexts)
        
        # Format prompt with specific instruction to use only provided context
        prompt = f"""Using *only* the context provided below, answer the question.  
If the answer cannot be determined from the context, say so.

Context:  
{context}

Question: {query_text}

Answer:"""

        # Get answer from LLM using chat API with error handling
        try:
            if not current_querying_model:
                raise ValueError("No querying model selected")
                
            response = ollama.chat(
                model=current_querying_model,
                messages=[{"role": "user", "content": prompt}]
            )
            
            logger.debug(f"Raw LLM response: {response}")
            
            # Extract content from the response, handling different response formats
            answer = None
            if hasattr(response, 'message'):
                # Object-style response
                if hasattr(response.message, 'content'):
                    answer = response.message.content
                elif isinstance(response.message, dict) and 'content' in response.message:
                    answer = response.message['content']
            elif isinstance(response, dict):
                # Dictionary-style response
                message = response.get('message')
                if isinstance(message, dict):
                    answer = message.get('content')
                elif hasattr(message, 'content'):
                    answer = message.content
            
            if not answer:
                logger.error(f"Could not extract answer from response: {response}")
                raise ValueError("No valid answer content in LLM response")
                
        except Exception as llm_error:
            error_message = f"Error getting answer from LLM: {str(llm_error)}"
            logger.error(error_message, exc_info=True)
            audio_path = text_to_speech(error_message)
            return {
                "type": "query_result",
                "success": False,
                "message": error_message,
                "answer": None,
                "sources": list(sources),
                "chunks_used": len(contexts),
                "audio_response_path": audio_path
            }
        
        # Generate audio response with retry logic
        audio_path = None
        max_retries = 3
        retry_delay = 1  # seconds
        
        for attempt in range(max_retries):
            try:
                audio_path = text_to_speech(answer)
                if audio_path and os.path.exists(audio_path) and os.path.getsize(audio_path) > 0:
                    break
                if attempt < max_retries - 1:
                    await asyncio.sleep(retry_delay)
            except Exception as e:
                logger.warning(f"TTS attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    await asyncio.sleep(retry_delay)
                else:
                    logger.error("All TTS attempts failed")
        
        result = {
            "type": "query_result",
            "success": True,
            "message": "Query processed successfully",
            "answer": answer,
            "sources": list(sources),
            "chunks_used": len(contexts),
            "audio_response_path": audio_path
        }
        
        # Log if audio wasn't generated
        if not audio_path:
            logger.warning("Audio response not available for successful query")
            
        return result
        
    except Exception as e:
        error_message = f"Error processing query: {str(e)}"
        try:
            audio_path = text_to_speech(error_message)
        except Exception as tts_error:
            logger.error(f"Failed to create audio for error message: {tts_error}")
            audio_path = None
            
        logger.error(f"Error processing query: {e}", exc_info=True)
        return {
            "type": "query_result",
            "success": False,
            "message": error_message,
            "answer": None,
            "sources": [],
            "chunks_used": 0,
            "audio_response_path": audio_path
        }
